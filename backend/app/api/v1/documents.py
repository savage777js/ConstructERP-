from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from sqlalchemy.orm import Session
from app.api import deps
from app.ai.service import ai_service
from app.models import core
import base64
import os
import uuid
from typing import Any

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    import docx
except ImportError:
    docx = None

router = APIRouter()

# Asegurar que el directorio de uploads existe
UPLOAD_DIR = "uploads/documents"
os.makedirs(UPLOAD_DIR, exist_ok=True)

def extract_text_from_docx(file_path: str) -> str:
    """Extrae texto de un archivo Word (.docx)."""
    if not docx:
        return "Error: python-docx no está instalado en el servidor."
    try:
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # También extraer texto de tablas
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = " ".join([p.text.strip() for p in cell.paragraphs if p.text.strip()])
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    full_text.append(" | ".join(row_text))
                    
        return "\n".join(full_text).strip()
    except Exception as e:
        return f"Error leyendo archivo Word (.docx): {str(e)}"

def extract_text_from_doc_fallback(file_path: str) -> str:
    """Intenta extraer texto legible de un archivo .doc binario como fallback."""
    try:
        with open(file_path, 'rb') as f:
            content = f.read()
        import re
        strings = re.findall(rb'[a-zA-Z0-9\s\.,;:!\?\-\(\)@_]{4,}', content)
        text_lines = []
        for s in strings:
            try:
                decoded = s.decode('utf-8', errors='ignore').strip()
                if len(decoded) > 5:
                    text_lines.append(decoded)
            except Exception:
                pass
        return "\n".join(text_lines)
    except Exception as e:
        return f"Error en fallback de lectura .doc: {str(e)}"

allow_ocr = deps.RoleChecker([
    core.UserRole.SUPER_ADMIN,
    core.UserRole.ADMIN,
    core.UserRole.HR_MANAGER,
    core.UserRole.PROJECT_MANAGER,
    core.UserRole.INVENTORY_MANAGER,
    core.UserRole.MANAGEMENT
])

@router.post("/ocr/invoice", dependencies=[Depends(allow_ocr)])
async def ocr_invoice(
    file: UploadFile = File(...),
    db: Session = Depends(deps.get_db),
    current_user: core.User = Depends(deps.get_current_user),
) -> Any:
    """Sube una factura (imagen, PDF o Word) y extrae sus datos usando IA."""
    
    # Validar extensión
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in [".jpg", ".jpeg", ".png", ".pdf", ".docx", ".doc", ".webp", ".jfif", ".heic", ".heif"]:
        raise HTTPException(status_code=400, detail="Formato de archivo no soportado. Use JPG, PNG, PDF, Word o formatos de imagen móviles (HEIC/WEBP/JFIF).")

    try:
        # Leer contenido
        contents = await file.read()
        
        # Guardar archivo localmente usando un nombre único seguro (previene Path Traversal)
        unique_filename = f"ocr_{uuid.uuid4()}{ext}"
        file_path = os.path.join(UPLOAD_DIR, unique_filename)
        with open(file_path, "wb") as f:
            f.write(contents)

        text_content = None
        image_base64 = None

        if ext == ".pdf":
            if not PdfReader:
                raise HTTPException(status_code=500, detail="El motor de lectura PDF (pypdf) no está instalado en el servidor.")
            
            reader = PdfReader(file_path)
            pages_text = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    pages_text.append(text)
            
            text_content = "\n".join(pages_text).strip()
            if not text_content:
                raise HTTPException(
                    status_code=400,
                    detail="El archivo PDF no contiene texto digital extraíble (podría ser un documento escaneado). Por favor, suba una versión con texto digital o una imagen JPG/PNG."
                )
        elif ext in [".docx", ".doc"]:
            if ext == ".docx":
                text_content = extract_text_from_docx(file_path)
            else:
                text_content = extract_text_from_doc_fallback(file_path)
                
            if not text_content or not text_content.strip():
                raise HTTPException(
                    status_code=400,
                    detail="No se pudo extraer texto legible del documento Word. Verifique que no esté vacío."
                )
        else:
            image_base64 = base64.b64encode(contents).decode("utf-8")
        
        # Llamar al servicio de IA
        ocr_data = await ai_service.process_document(
            image_base64=image_base64,
            text_content=text_content,
            category="invoice"
        )
        
        if not ocr_data:
            raise HTTPException(status_code=500, detail="La IA no pudo procesar el documento.")

        return {
            "filename": file.filename,
            "data": ocr_data,
            "preview_url": f"/uploads/documents/{unique_filename}"
        }

    except HTTPException:
        raise
    except Exception as e:
        print(f"Error en OCR: {e}")
        raise HTTPException(status_code=500, detail=str(e))

def create_docx_from_ocr(title: str, ocr_data: dict, output_path: str):
    """Crea un documento de Word (.docx) a partir de los datos del OCR."""
    if not docx:
        raise Exception("python-docx no está instalado en el servidor.")
    
    doc = docx.Document()
    doc.add_heading(f"Reporte de Extracción OCR - {title}", level=0)
    doc.add_paragraph("Generado automáticamente por ConstructERP AI\n")
    
    # Agregar tabla de datos estructurados
    doc.add_heading("Datos Estructurados Extraídos", level=1)
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Shading Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Campo'
    hdr_cells[1].text = 'Valor'
    
    for k, v in ocr_data.items():
        if k == 'confidence':
            continue
        row_cells = table.add_row().cells
        key_label = k.replace('_', ' ').title()
        row_cells[0].text = str(key_label)
        row_cells[1].text = str(v) if v is not None else "No detectado"
        
    doc.add_paragraph("\n")
    
    # Si hay un resumen, descripción o detalle principal
    doc.add_heading("Resumen de Contenido", level=1)
    resumen_text = ocr_data.get("resumen") or ocr_data.get("description") or ocr_data.get("detalles")
    if resumen_text:
        doc.add_paragraph(str(resumen_text))
    else:
        doc.add_paragraph("Datos extraídos con éxito según el tipo de documento.")
        
    doc.save(output_path)

@router.post("/")
async def upload_document(
    employee_id: int = Form(...),
    category: str = Form(...),
    title: str = Form(...),
    file: UploadFile = File(...),
    db: Session = Depends(deps.get_db),
    current_user: core.User = Depends(deps.get_current_user),
):
    """Sube un documento para un trabajador."""
    # Validar que el trabajador existe
    employee = db.query(core.Employee).filter(core.Employee.id == employee_id).first()
    if not employee:
        raise HTTPException(status_code=404, detail="Trabajador no encontrado")

    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in [".jpg", ".jpeg", ".png", ".pdf", ".docx", ".doc", ".webp", ".jfif", ".heic", ".heif"]:
        raise HTTPException(status_code=400, detail="Formato de archivo no soportado. Use JPG, PNG, PDF, Word o imágenes móviles (HEIC/WEBP/JFIF).")

    contents = await file.read()
    
    # Crear un nombre único de archivo original
    unique_filename = f"{uuid.uuid4()}{ext}"
    file_path = os.path.join(UPLOAD_DIR, unique_filename)
    
    with open(file_path, "wb") as f:
        f.write(contents)

    # Si es una organización, asociamos a la misma org del usuario
    org_id = current_user.organization_id if current_user.organization_id else employee.organization_id

    file_type = file.content_type
    file_size = len(contents)
    ocr_status = "PENDING"
    ocr_content = None
    extracted_data = None

    if ext in [".jpg", ".jpeg", ".png", ".webp", ".jfif", ".heic", ".heif"]:
        import json
        image_base64 = base64.b64encode(contents).decode("utf-8")
        try:
            ocr_data = await ai_service.process_document(
                image_base64=image_base64,
                text_content=None,
                category=category
            )
        except Exception:
            ocr_data = None
            
        if not ocr_data:
            ocr_data = {
                "tipo_documento": category,
                "titulo": title,
                "resumen": "El procesamiento OCR no retornó datos estructurados válidos.",
                "confidence": 0.0
            }
            
        docx_filename = f"{uuid.uuid4()}.docx"
        docx_path = os.path.join(UPLOAD_DIR, docx_filename)
        
        try:
            create_docx_from_ocr(title, ocr_data, docx_path)
            unique_filename = docx_filename
            file_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            file_size = os.path.getsize(docx_path)
            ocr_status = "COMPLETED"
            ocr_content = json.dumps(ocr_data, indent=2, ensure_ascii=False)
            extracted_data = ocr_data
        except Exception as e:
            print(f"Error generando Word de OCR: {e}")

    # Si es un documento obligatorio (Contrato, Licencia, Cédula, Certificado), reemplazamos el anterior para no duplicar
    if category in ["Contrato", "Licencia", "Cédula", "Certificado"]:
        existing_docs = db.query(core.Document).filter(
            core.Document.employee_id == employee_id,
            core.Document.category == category
        ).all()
        for old_doc in existing_docs:
            try:
                local_path = old_doc.file_path.lstrip('/')
                if os.path.exists(local_path):
                    os.remove(local_path)
            except Exception as e:
                print(f"Error borrando archivo antiguo ({category}): {e}")
            db.delete(old_doc)
        db.commit()

    db_doc = core.Document(
        organization_id=org_id,
        title=title,
        file_path=f"/uploads/documents/{unique_filename}",
        file_type=file_type,
        file_size=file_size,
        category=category,
        employee_id=employee_id,
        created_by=current_user.id,
        ocr_status=ocr_status,
        ocr_content=ocr_content,
        extracted_data=extracted_data
    )
    db.add(db_doc)
    db.commit()
    db.refresh(db_doc)

    return {
        "id": str(db_doc.id),
        "title": db_doc.title,
        "file_path": db_doc.file_path,
        "category": db_doc.category,
        "ocr_status": db_doc.ocr_status,
        "created_at": db_doc.created_at
    }

@router.get("/employee/{employee_id}")
def list_employee_documents(
    employee_id: int,
    db: Session = Depends(deps.get_db),
    current_user: core.User = Depends(deps.get_current_user),
):
    """Lista todos los documentos cargados para un trabajador específico."""
    # Verificar pertenencia a la organización
    query_emp = db.query(core.Employee).filter(core.Employee.id == employee_id)
    if current_user.organization_id:
        query_emp = query_emp.filter(core.Employee.organization_id == current_user.organization_id)
    employee = query_emp.first()
    if not employee:
        raise HTTPException(status_code=404, detail="Trabajador no encontrado en su organización")

    docs = db.query(core.Document).filter(core.Document.employee_id == employee_id).all()
    return [{
        "id": str(d.id),
        "title": d.title,
        "file_path": d.file_path,
        "category": d.category,
        "ocr_status": d.ocr_status,
        "ocr_content": d.ocr_content,
        "extracted_data": d.extracted_data,
        "created_at": d.created_at
    } for d in docs]

@router.get("/employee/{employee_id}/download-zip")
def download_employee_documents_zip(
    employee_id: int,
    db: Session = Depends(deps.get_db),
    current_user: core.User = Depends(deps.get_current_user),
):
    """Genera y descarga un archivo ZIP con todos los documentos del trabajador."""
    import zipfile
    import io
    from fastapi.responses import Response

    query_emp = db.query(core.Employee).filter(core.Employee.id == employee_id)
    if current_user.organization_id:
        query_emp = query_emp.filter(core.Employee.organization_id == current_user.organization_id)
    employee = query_emp.first()
    if not employee:
        raise HTTPException(status_code=404, detail="Trabajador no encontrado en su organización")
        
    docs = db.query(core.Document).filter(core.Document.employee_id == employee_id).all()
    if not docs:
        raise HTTPException(status_code=404, detail="El trabajador no tiene documentos en su carpeta digital")
        
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
        for doc in docs:
            local_path = doc.file_path.lstrip('/')
            doc.ensure_local_file(db)
            if os.path.exists(local_path):
                ext = os.path.splitext(local_path)[1]
                safe_title = "".join([c for c in doc.title if c.isalnum() or c in [' ', '_', '-']]).strip()
                if not safe_title:
                    safe_title = f"documento_{doc.id}"
                filename_in_zip = f"{safe_title}{ext}"
                zip_file.write(local_path, filename_in_zip)
            else:
                print(f"Warning: local file {local_path} not found for ZIP.")
                
    zip_buffer.seek(0)
    zip_name = f"documentos_{employee.first_name}_{employee.last_name}.zip".replace(" ", "_")
    
    return Response(
        content=zip_buffer.getvalue(),
        media_type="application/x-zip-compressed",
        headers={
            "Content-Disposition": f"attachment; filename={zip_name}"
        }
    )

@router.delete("/{document_id}")
def delete_document(
    document_id: str,
    db: Session = Depends(deps.get_db),
    current_user: core.User = Depends(deps.get_current_user),
):
    """Elimina un documento."""
    query = db.query(core.Document).filter(core.Document.id == document_id)
    if current_user.organization_id:
        query = query.filter(core.Document.organization_id == current_user.organization_id)
    doc = query.first()
    if not doc:
        raise HTTPException(status_code=404, detail="Documento no encontrado o no pertenece a su organización")
        
    db.delete(doc)
    db.commit()
    return {"message": "Documento eliminado correctamente"}

@router.post("/{document_id}/ocr", dependencies=[Depends(allow_ocr)])
async def process_document_ocr(
    document_id: str,
    db: Session = Depends(deps.get_db),
    current_user: core.User = Depends(deps.get_current_user),
):
    """Procesa mediante OCR un documento ya existente en la base de datos (imagen, PDF o Word)."""
    import json

    query = db.query(core.Document).filter(core.Document.id == document_id)
    if current_user.organization_id:
        query = query.filter(core.Document.organization_id == current_user.organization_id)
    doc = query.first()
    if not doc:
        raise HTTPException(status_code=404, detail="Documento no encontrado o no pertenece a su organización")
        
    # Verificar que el archivo sea una imagen, PDF o Word
    ext = os.path.splitext(doc.file_path)[1].lower()
    if ext not in [".jpg", ".jpeg", ".png", ".pdf", ".docx", ".doc"]:
        raise HTTPException(status_code=400, detail="El procesamiento OCR mediante IA solo es soportado para imágenes, archivos PDF y Word (DOCX/DOC).")

    # Recuperar archivo local (auto-curación desde BD si fue borrado por Render)
    local_path = doc.file_path.lstrip('/')
    if not doc.ensure_local_file(db):
        raise HTTPException(status_code=404, detail="El archivo físico del documento no existe en el servidor ni en la base de datos.")

    text_content = None
    image_base64 = None

    try:
        if ext == ".pdf":
            if not PdfReader:
                raise HTTPException(status_code=500, detail="El motor de lectura PDF (pypdf) no está instalado en el servidor.")
            
            reader = PdfReader(local_path)
            pages_text = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    pages_text.append(text)
            
            text_content = "\n".join(pages_text).strip()
            if not text_content:
                raise HTTPException(
                    status_code=400,
                    detail="El archivo PDF no contiene texto digital extraíble (podría ser un documento escaneado). Por favor, suba una versión con texto digital o una imagen JPG/PNG."
                )
        elif ext in [".docx", ".doc"]:
            if ext == ".docx":
                text_content = extract_text_from_docx(local_path)
            else:
                text_content = extract_text_from_doc_fallback(local_path)
            
            if not text_content or not text_content.strip():
                raise HTTPException(
                    status_code=400,
                    detail="No se pudo extraer texto legible del documento Word. Verifique que no esté vacío."
                )
        else:
            with open(local_path, "rb") as f:
                contents = f.read()
            image_base64 = base64.b64encode(contents).decode("utf-8")

        # Invocar servicio de IA con el contenido del documento y su categoría
        ocr_data = await ai_service.process_document(
            image_base64=image_base64,
            text_content=text_content,
            category=doc.category
        )

        if ocr_data:
            doc.ocr_status = "COMPLETED"
            # Guardamos el JSON formateado de forma legible en ocr_content
            doc.ocr_content = json.dumps(ocr_data, indent=2, ensure_ascii=False)
            doc.extracted_data = ocr_data
            db.commit()
            db.refresh(doc)
            return {
                "id": str(doc.id),
                "ocr_status": doc.ocr_status,
                "data": ocr_data
            }
        else:
            doc.ocr_status = "FAILED"
            db.commit()
            raise HTTPException(status_code=500, detail="La IA no pudo extraer datos del documento.")
            
    except HTTPException:
        raise
    except Exception as e:
        doc.ocr_status = "FAILED"
        db.commit()
        raise HTTPException(status_code=500, detail=str(e))
